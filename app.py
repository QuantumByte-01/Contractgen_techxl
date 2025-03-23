import os
import io
import docx
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename

from utils import (
    parse_clauses,
    generate_contract,
    remove_suggestions_from_html,
    extract_suggestions_from_html,
    incorporate_suggestions_with_model,
)
from chat import chat_with_model
from analysis import extract_text_from_pdf, extract_text_from_docx, summarize_and_analyze

app = Flask(__name__)
app.secret_key = "your_secret_key_here"  # For demonstration, but we won't store the contract in session.

# For analyzing existing PDF/DOCX
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {"pdf", "docx"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/")
def index():
    # Main page with links to contract forms & analysis
    return render_template("index.html")

@app.route("/form")
def form():
    contract_type = request.args.get("contract_type", "nda")
    if contract_type == "service_agreement":
        return render_template("form_service.html")
    elif contract_type == "employment_contract":
        return render_template("form_employment.html")
    elif contract_type == "rental_agreement":
        return render_template("form_rental.html")
    else:
        return render_template("form_nda.html")

@app.route("/generate", methods=["POST"])
def generate():
    contract_type = request.form.get("contract_type", "nda")
    other_clauses = {}

    if contract_type == "nda":
        details = {
            "EFFECTIVE_DATE": request.form.get("EFFECTIVE_DATE", "2025-01-01"),
            "PARTY_A": request.form.get("PARTY_A", "Party A"),
            "PARTY_B": request.form.get("PARTY_B", "Party B")
        }
        customization = {
            "JURISDICTION": request.form.get("JURISDICTION", "Default Jurisdiction"),
            "TONE": request.form.get("TONE", "formal")
        }
        clauses = {
            "Confidentiality Clause": request.form.get("CLAUSE_CONFIDENTIALITY", ""),
            "Non-Use Clause": request.form.get("CLAUSE_NON_USE", ""),
            "Term Clause": request.form.get("CLAUSE_TERM", "")
        }
        other_input = request.form.get("OTHER_CLAUSES", "")
        if other_input.strip():
            other_clauses = parse_clauses(other_input)

    elif contract_type == "service_agreement":
        details = {
            "EFFECTIVE_DATE": request.form.get("EFFECTIVE_DATE", "2025-01-01"),
            "SERVICE_PROVIDER": request.form.get("SERVICE_PROVIDER", "Service Provider"),
            "CLIENT": request.form.get("CLIENT", "Client")
        }
        customization = {
            "JURISDICTION": request.form.get("JURISDICTION", "Default Jurisdiction"),
            "TONE": request.form.get("TONE", "formal")
        }
        clauses = {
            "Scope of Services": request.form.get("CLAUSE_SCOPE_OF_SERVICES", ""),
            "Payment Terms": request.form.get("CLAUSE_PAYMENT_TERMS", ""),
            "Confidentiality Clause": request.form.get("CLAUSE_CONFIDENTIALITY", ""),
            "Non-Compete Clause": request.form.get("CLAUSE_NON_COMPETE", ""),
            "Intellectual Property Clause": request.form.get("CLAUSE_INTELLECTUAL_PROPERTY", ""),
            "Indemnification Clause": request.form.get("CLAUSE_INDEMNIFICATION", ""),
            "Termination Clause": request.form.get("CLAUSE_TERMINATION", ""),
            "Dispute Resolution Clause": request.form.get("CLAUSE_DISPUTE_RESOLUTION", "")
        }
        other_input = request.form.get("OTHER_CLAUSES", "")
        if other_input.strip():
            other_clauses = parse_clauses(other_input)

    elif contract_type == "employment_contract":
        details = {
            "EFFECTIVE_DATE": request.form.get("EFFECTIVE_DATE", "2025-01-01"),
            "EMPLOYER": request.form.get("EMPLOYER", "Employer"),
            "EMPLOYEE": request.form.get("EMPLOYEE", "Employee")
        }
        customization = {
            "JURISDICTION": request.form.get("JURISDICTION", "Default Jurisdiction"),
            "TONE": request.form.get("TONE", "formal")
        }
        clauses = {
            "Job Description": request.form.get("CLAUSE_JOB_DESCRIPTION", ""),
            "Compensation": request.form.get("CLAUSE_COMPENSATION", ""),
            "Benefits": request.form.get("CLAUSE_BENEFITS", ""),
            "Termination Clause": request.form.get("CLAUSE_TERMINATION", ""),
            "Confidentiality Clause": request.form.get("CLAUSE_CONFIDENTIALITY", ""),
            "Non-Compete Clause": request.form.get("CLAUSE_NON_COMPETE", "")
        }
        other_input = request.form.get("OTHER_CLAUSES", "")
        if other_input.strip():
            other_clauses = parse_clauses(other_input)

    elif contract_type == "rental_agreement":
        details = {
            "EFFECTIVE_DATE": request.form.get("EFFECTIVE_DATE", "2025-01-01"),
            "LANDLORD": request.form.get("LANDLORD", "Landlord"),
            "TENANT": request.form.get("TENANT", "Tenant")
        }
        customization = {
            "JURISDICTION": request.form.get("JURISDICTION", "Default Jurisdiction"),
            "TONE": request.form.get("TONE", "formal")
        }
        clauses = {
            "Premises Description": request.form.get("CLAUSE_PREMISES_DESCRIPTION", ""),
            "Rent Payment": request.form.get("CLAUSE_RENT_PAYMENT", ""),
            "Security Deposit": request.form.get("CLAUSE_SECURITY_DEPOSIT", ""),
            "Termination Clause": request.form.get("CLAUSE_TERMINATION", ""),
            "Maintenance Responsibility": request.form.get("CLAUSE_MAINTENANCE_RESPONSIBILITY", "")
        }
        other_input = request.form.get("OTHER_CLAUSES", "")
        if other_input.strip():
            other_clauses = parse_clauses(other_input)
    else:
        return "Unsupported contract type."

    # Generate final contract HTML
    contract_html = generate_contract(contract_type, details, clauses, other_clauses, customization)
    # We do NOT store it in session. We'll rely on localStorage in the client.

    return render_template("result.html", contract=contract_html)

@app.route("/apply", methods=["POST"])
def apply_suggestions():
    """
    1) Extract the contract from localStorage (hidden field).
    2) Separate the main contract from the suggestions block.
    3) Feed both to the model, asking it to produce a final version with NO suggestions.
    4) Return that final version in result.html.
    """
    contract_html = request.form.get("contract_html", "No contract generated.")
    if contract_html == "No contract generated.":
        return render_template("result.html", contract=contract_html)

    # 1. Separate main contract from suggestions
    main_body, suggestions_block = extract_suggestions_from_html(contract_html)

    if not suggestions_block:
        # If no suggestions exist, just return the same contract
        final_version = main_body
    else:
        # 2. Incorporate suggestions via the model
        final_version = incorporate_suggestions_with_model(main_body, suggestions_block)

    return render_template("result.html", contract=final_version)


@app.route("/chat", methods=["GET", "POST"])
def chat():
    """
    Allows iterative modification of the contract. The contract is loaded from localStorage
    and sent to the server with the user's request.
    """
    if request.method == "GET":
        # Just show a page with a text area for user input
        return render_template("chat.html", contract="")
    else:
        # POST: user has typed a message and provided contract_html from localStorage
        contract_html = request.form.get("contract_html", "No contract generated.")
        user_message = request.form.get("user_message", "")
        if contract_html == "No contract generated.":
            return render_template("chat.html", contract=contract_html)

        updated_contract = chat_with_model(contract_html, user_message)
        return render_template("chat.html", contract=updated_contract)

@app.route("/export", methods=["POST"])
def export_docx():
    """
    Exports the contract as a DOCX file, removing HTML tags and suggestions/analysis.
    """
    contract_html = request.form.get("contract_html", "No contract generated.")
    if contract_html == "No contract generated.":
        return "No contract generated."

    # 1) Remove suggestions block
    contract_no_suggestions = remove_suggestions_from_html(contract_html)
    # 2) Convert HTML -> plain text for docx
    plain_text = html_to_text(contract_no_suggestions)

    # 3) Build docx
    doc = docx.Document()
    doc.add_heading("Exported Contract", 0)
    doc.add_paragraph(plain_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="contract.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.route("/analyze", methods=["GET", "POST"])
def analyze():
    """
    Summarize & analyze an existing PDF or DOCX.
    """
    if request.method == "GET":
        return render_template("analysis.html")
    else:
        if "file" not in request.files:
            return "No file part."
        file = request.files["file"]
        if file.filename == "":
            return "No file selected."
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(filepath)

            # Extract text
            if filename.lower().endswith(".pdf"):
                text_content = extract_text_from_pdf(filepath)
            else:
                text_content = extract_text_from_docx(filepath)

            analysis_result = summarize_and_analyze(text_content)
            return render_template("analysis_result.html", analysis=analysis_result)
        else:
            return "Unsupported file format."
        
import re

def html_to_text(html_content: str) -> str:
    """
    A simple approach to remove HTML tags, leaving plain text.
    For more robust solutions, consider 'BeautifulSoup' or 'html2text'.
    """
    # Remove tags
    text = re.sub(r'<[^>]*>', '', html_content)
    # Unescape common entities if needed
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&')
    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text


if __name__ == "__main__":
    app.run(debug=True)
